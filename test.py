generator.py
import uuid
from datetime import datetime, timezone
import os
from dotenv import load_dotenv
from backend.models.company_profile import CompanyProfile
from backend.prompts.loader import build_section_prompt, load_prompt
from backend.generation.validator import validate_draft_llm
from backend.prompts.type_behavior import get_type_behavior, should_generate_toc, get_forbidden_phrases
from backend.prompts.risk_behavior import get_risk_behavior
from backend.prompts.section_rules import get_section_rules, get_section_word_limit
from langchain_core.prompts import ChatPromptTemplate
from backend.generation.llm_provider import get_llm
from langchain_core.messages import SystemMessage, HumanMessage
import json
import re

llm = get_llm()

load_dotenv()


def _should_generate_section(doc_type: str, section_name: str) -> bool:
    """
    Returns False for sections that must be skipped for this doc type.
    Currently gates:
      - Table of Contents / Index  →  only for POLICY, SOP, REPORT, HANDBOOK, STRATEGY, PROPOSAL
    All other sections always return True.
    """
    section_lower = section_name.lower()

    toc_keywords = ["table of contents", "index", "contents page"]
    is_toc_section = any(kw in section_lower for kw in toc_keywords)

    if is_toc_section:
        return should_generate_toc(doc_type)

    return True



# SECTION VALIDATOR
# Checks the LLM output for common quality issues.

def _validate_section_output(
    content: str,
    section_name: str,
    doc_type: str
) -> dict:
    """
    Validates a single generated section.

    Returns:
        {
            "valid": bool,
            "issues": list[str],
            "word_count": int,
            "min_words": int,
            "max_words": int
        }
    """
    issues = []
    word_count = len(content.split())
    min_words, max_words = get_section_word_limit(doc_type, section_name)

    repetitive_phrases = [
        "this section constitutes a binding policy requirement",
        "all employees are subject to this policy from their start date",
        "violations of this policy may result in disciplinary action",
        "this policy is reviewed annually"
    ]

    for phrase in repetitive_phrases:
        if phrase in content.lower():
            issues.append(f"Repetitive boilerplate detected: '{phrase}'")
    
    instruction_phrases = [
    "enter the",
    "provide the",
    "complete all fields",
    "insert the",
    "fill in",
    "specify the",
    "record the following"
    ]

    for phrase in instruction_phrases:
        if phrase in content.lower() and doc_type not in ["FORM", "TEMPLATE"]:
            issues.append(f"Instructional language not allowed in {doc_type}: '{phrase}'")


    #  Word count checks 
    if word_count < min_words:
        issues.append(
            f"Too short: {word_count} words (minimum required: {min_words})"
        )

    if word_count > max_words:
        issues.append(
            f"Too long: {word_count} words (maximum allowed: {max_words})"
        )

    #  Section heading repeated inside content body 
    if section_name.lower() in content[:200].lower():
        issues.append(
            "Section heading is repeated inside the content body. "
            "Heading is added externally — remove it from the content."
        )

    #  Placeholder / unfilled text 
    bad_placeholders = [
        "[TO BE FILLED]", "[INSERT HERE]", "[TBD]",
        "TODO:", "[PLACEHOLDER]", "[ADD CONTENT HERE]",
        "[COMPANY NAME]", "[DATE]"
    ]
    for ph in bad_placeholders:
        if ph.upper() in content.upper():
            issues.append(f"Unfilled placeholder found: '{ph}'")

    #  Model preamble leak 
    preamble_phrases = [
        "here is the content",
        "here is the section",
        "below is the content",
        "i'll now generate",
        "the following section",
        "as requested, here"
    ]
    content_start = content[:120].lower()
    for phrase in preamble_phrases:
        if phrase in content_start:
            issues.append(
                f"Model added preamble text: '{phrase}'. "
                "Output must start directly with document content."
            )

    #  Forbidden phrases 
    forbidden = get_forbidden_phrases(doc_type)
    for phrase in forbidden:
        if phrase.lower() in content.lower():
            issues.append(f"Forbidden phrase detected: '{phrase}'")

    #  Empty output 
    if not content.strip():
        issues.append("Generated content is empty.")

    return {
        "valid": len(issues) == 0,
        "issues": issues,
        "word_count": word_count,
        "min_words": min_words,
        "max_words": max_words
    }


# SINGLE SECTION GENERATOR
# Calls AzureOpenAI for one section and validates output.

def _generate_single_section(
    section_name: str,
    mandatory: bool,
    registry_doc: dict,
    company_block: str,
    company_profile: dict,
    inputs_block: str,
    industry_context: str,
    user_notes: str,
    all_sections: list,
    retry: bool = False,
    previous_issues: list = None
) -> dict:
    """
    Generates content for one section.
    Validates the output and returns the full section result.

    Returns:
        {
            "name": str,
            "mandatory": bool,
            "content": str,
            "section_validation": dict   ← NEW: per-section quality result
        }
    """
    doc_type      = registry_doc["internal_type"]
    risk_level    = registry_doc["risk_level"]
    type_behavior_data  = get_type_behavior(doc_type)
    tone = type_behavior_data.get("tone", "professional")
    voice = type_behavior_data.get("voice", "third-person")
    format_style = type_behavior_data.get("format", "")
    rules = type_behavior_data.get("rules", "")
    avg_section_words = type_behavior_data.get("avg_section_words", "")
    risk_behavior  = get_risk_behavior(risk_level)
    section_rules  = get_section_rules(doc_type, section_name)
    company_name = company_profile.get("company_name", "") if company_profile else ""
    industry = company_profile.get("industry", "") if company_profile else ""
    employee_count = company_profile.get("employee_count", "") if company_profile else ""
    region = ", ".join(company_profile.get("regions", [])) if company_profile else ""
    jurisdiction = company_profile.get("default_jurisdiction", "") if company_profile else ""

    # Build TOC section list string for the prompt
    all_sections_str = "\n".join(
        f"{i+1}. {s['name']}"
        for i, s in enumerate(all_sections)
    )
    min_words, max_words = get_section_word_limit(doc_type, section_name)
    if max_words > 300:
        max_words = 300

    forbidden_phrases = get_forbidden_phrases(doc_type)

    context = {
        "document_name":   registry_doc["document_name"],
        "document_type":   doc_type,
        "risk_level":      risk_level,
        "section_name":    section_name,
        "mandatory":       str(mandatory),
        "company_profile": company_profile,
        "document_inputs": inputs_block,
        "industry_context": industry_context,
        "type_behavior":   rules,
        "tone": tone,
        "voice": voice,
        "format_style": format_style,
        "avg_section_words": avg_section_words,
        "risk_behavior":   risk_behavior,
        "section_rules":   section_rules,
        "all_sections":    all_sections_str,
        "toc_required":    str(should_generate_toc(doc_type)).upper(),
        "min_words": min_words,
        "max_words": max_words,
        "company_name": company_name,
        "industry": industry,
        "employee_count": employee_count,
        "region": region,
        "jurisdiction": jurisdiction,
        "forbidden_phrases": "\n".join(forbidden_phrases)
    }

    base_prompt = build_section_prompt(context)

    # ── Add retry context if this is a re-generation ───────
    retry_block = ""
    if retry and previous_issues:
        retry_block = (
            "\n\nPREVIOUS ATTEMPT FAILED VALIDATION — FIX ALL ISSUES BELOW:\n"
            + "\n".join(f"  • {issue}" for issue in previous_issues)
            + "\n\nRe-generate the section addressing every issue listed above.\n"
        )

    full_prompt = f"""
{base_prompt}
{retry_block}
Additional Notes:
{user_notes or "None provided."}
""".strip()

    system_message = f"""
    You are generating the FINAL VERSION of an enterprise {doc_type} document.

    You are NOT:
    - Writing instructions
    - Writing guidance
    - Writing meta commentary
    - Writing a template unless document_type == TEMPLATE
    - Writing placeholders unless document_type == FORM or TEMPLATE

    You ARE:
    - Writing the actual content as it will appear in the published document.

    STRICT LENGTH RULE:
    - Between {min_words} and {max_words} words.
    - Do NOT exceed limit.
    - If exceeded, you FAIL.

    STRICT OUTPUT RULES:
    - Start directly with content.
    - Do NOT repeat section title.
    - Do NOT explain what to do.
    - Do NOT include examples unless explicitly required.
    - Do NOT add filler language.

    SECTION CONTEXT CONTROL:
    - Write content ONLY relevant to the section name.
    - Do NOT introduce topics that belong to other enterprise policies.
    - Do NOT expand scope beyond the purpose of this specific document.

    """

    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=full_prompt)
    ]


    response = llm.invoke(messages)

    try:
        content = getattr(response, "content", str(response)).strip()
    except:
        content = str(response).strip()

    structured_content = None

    # Clean possible ```json blocks
    cleaned = content.strip()

    # Remove markdown code block if present
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z]*", "", cleaned)
        cleaned = cleaned.replace("```", "").strip()

    try:
        parsed = json.loads(cleaned)

        # CASE 1: direct table block
        if isinstance(parsed, dict) and "table" in parsed:
            structured_content = parsed
            content = ""

        # CASE 2: list of mixed blocks
        elif isinstance(parsed, list):
            structured_content = parsed
            content = ""

    except Exception as e:
        structured_content = None

    max_words_allowed = max_words
    words = content.split()
    if len(words) > max_words_allowed:
        content = " ".join(words[:max_words_allowed])

    # Validate output 
    validation_target = content

    if structured_content:
        validation_target = json.dumps(structured_content)

    section_validation = _validate_section_output(
        content=validation_target,
        section_name=section_name,
        doc_type=doc_type
    )
    print("LLM RAW RESULT:", response)
    print("LLM CONTENT:", response.content)

    return {
        "name":               section_name,
        "mandatory":          mandatory,
        "content":            content,
        "structured_content": structured_content,
        "section_validation": section_validation
    }


# SECTION REGENERATION (User-triggered from UI)

def regenerate_section_llm(draft: dict, section: dict, issues: list) -> str:

    template = load_prompt("regeneration_prompt")

    formatted_prompt = template.format(
        document_type=draft["source_document"]["internal_type"],
        risk_level=draft["source_document"]["risk_level"],
        department=draft["source_document"]["department"],
        section_name=section["name"],
        original_content=section["content"],
        issues="\n".join(issues)
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert enterprise document improver."),
        ("human", formatted_prompt)
    ])

    chain = prompt | llm
    result = chain.invoke({})

    return result.content.strip()

def _compress_sections(sections, max_words):
    per_section_budget = max_words // len(sections)

    for s in sections:
        words = s["content"].split()
        if len(words) > per_section_budget:
            s["content"] = " ".join(words[:per_section_budget])

    return sections

# MAIN GENERATE DRAFT

def generate_draft(
    registry_doc: dict,
    department: str,
    document_filename: str,
    company_profile: CompanyProfile = None,
    document_inputs: dict = None,
    user_notes: str = None
) -> dict:
    """
    Generates a full document draft section by section.

    Flow:
      1. Build draft skeleton
      2. Format company profile + user inputs
      3. For each section:
            a. TOC gate  → skip if not needed for this doc type
            b. Generate  → call LLM
            c. Validate  → check word count, placeholders, preamble, etc.
            d. Auto-retry once if validation fails
      4. Run full-draft AI validation with regeneration loop
      5. Return final draft

    Returns: draft dict (same shape as before + section_validation per section)
    """

    #  Step 1: Draft  
    draft = {
        "draft_id": str(uuid.uuid4()),
        "source_document": {
            "department":           department,
            "document_filename":    document_filename,
            "document_name":        registry_doc["document_name"],
            "internal_type":        registry_doc["internal_type"],
            "risk_level":           registry_doc["risk_level"],
            "compliance_alignment": registry_doc.get("compliance_alignment", [])
        },
        "version": "v1.0",
        "status": "DRAFT",
        "generation_metadata": {
            "generated_at":    datetime.now(timezone.utc).isoformat(),
            "generated_by":    "azure_openai",
            "deterministic":   True,
            "prompt_version":  "v2",           # ← bumped to v2 after enhancement
            "toc_generated":   should_generate_toc(registry_doc["internal_type"]),
            "retry_count":     0
        },
        "sections": [],
        "validation": {
            "status": "NOT_RUN",
            "issues": []
        },
        "approval": {
            "required":     registry_doc["approval_required"],
            "approved":     False,
            "approved_by":  None,
            "approved_at":  None
        }
    }

    #  Step 2: Format context blocks 
    company_block = ""
    if company_profile:
        company_block = (
            f"Company Name: {company_profile.get('company_name')}\n"
            f"Industry: {company_profile.get('industry')}\n"
            f"Employee Count: {company_profile.get('employee_count')}\n"
            f"Region: {', '.join(company_profile.get('regions', []))}\n"
            f"Compliance: {', '.join(company_profile.get('compliance_frameworks', []))}\n"
            f"Jurisdiction: {company_profile.get('default_jurisdiction')}\n"
        )

    inputs_block = ""
    if document_inputs:
        for key, value in document_inputs.items():
            inputs_block += f"{key}: {value}\n"


    all_sections     = registry_doc["sections"]

    #  Step 3: Generate each section 
    SECTION_MAX_RETRIES = 1   # One auto-retry per section

    industry_context = load_prompt("industry_context")

    for section in all_sections:
        section_name = section["name"]
        mandatory    = section["mandatory"]

        if section_name.lower() in [
            "security",
            "compliance",
            "data protection",
            "incident response"
        ]: 
            industry_block = industry_context
        else:
            industry_block = ""  


        if not _should_generate_section(
            registry_doc["internal_type"], section_name
        ):
            print(f"[SKIP] '{section_name}' — not required for {registry_doc['internal_type']}")
            continue

        print(f"[GEN]  Generating section: '{section_name}'")

        #  First attempt 
        section_result = _generate_single_section(
            section_name=section_name,
            mandatory=mandatory,
            registry_doc=registry_doc,
            company_profile=company_profile,
            company_block=company_block,
            inputs_block=inputs_block,
            industry_context=industry_block,
            user_notes=user_notes,
            all_sections=all_sections,
            retry  =False
        )

        #  Section-level auto-retry 
        if not section_result["section_validation"]["valid"]:
            issues = section_result["section_validation"]["issues"]
            print(
                f"[WARN] Section '{section_name}' failed validation "
                f"({len(issues)} issue(s)). Retrying..."
            )

            retry_result = _generate_single_section(
                section_name=section_name,
                mandatory=mandatory,
                registry_doc=registry_doc,
                company_block=company_block,
                company_profile=company_profile,    
                inputs_block=inputs_block,
                industry_context=industry_block,
                user_notes=user_notes,
                all_sections=all_sections,
                retry=True,
                previous_issues=issues
            )

            # Use retry result only if it's better (or equal)
            if (
                retry_result["section_validation"]["valid"]
                or len(retry_result["section_validation"]["issues"])
                <= len(issues)
            ):
                section_result = retry_result

        draft["sections"].append(section_result)
        print(
            f"[DONE] '{section_name}' — "
            f"{section_result['section_validation']['word_count']} words | "
            f"valid: {section_result['section_validation']['valid']}"
        )

        #  Global document word cap (max ~2000 words ≈ 4 pages) 

        MAX_TOTAL_WORDS = 1800
        total_words = sum(
            s["section_validation"]["word_count"]
            for s in draft["sections"]
        )

        print(f"[INFO] Total document words before trim: {total_words}")

        if total_words > MAX_TOTAL_WORDS:
            draft["sections"] = _compress_sections(draft["sections"], MAX_TOTAL_WORDS)


    #  Step 4: Full-draft AI validation + regeneration loop 
    MAX_DRAFT_RETRIES = 2
    retry_count       = 0

    validation_result = {"status": "NOT_RUN", "issues": []}

    while retry_count <= MAX_DRAFT_RETRIES:

        try:
            validation_result = validate_draft_llm(draft)
        except Exception as e:
            validation_result = {
                "status": "ERROR",
                "issues": [f"Validation failed: {str(e)}"]
            }

        print(f"[VALIDATE] Status: {validation_result['status']} | Retry: {retry_count}")

        draft["validation"]                            = validation_result
        draft["generation_metadata"]["retry_count"]    = retry_count

        if validation_result["status"] == "PASS":
            draft["status"] = "READY_FOR_APPROVAL"
            break
        else:
            draft["status"] = "NEEDS_REVIEW"

        if retry_count < MAX_DRAFT_RETRIES:
            issues = validation_result.get("issues", [])

            for section in draft["sections"]:
                if section["mandatory"]:
                    try:
                        improved = regenerate_section_llm(
                            draft=draft,
                            section=section,
                            issues=issues
                        )
                        section["content"] = improved

                        # Re-validate the regenerated section
                        section["section_validation"] = _validate_section_output(
                            content=improved,
                            section_name=section["name"],
                            doc_type=registry_doc["internal_type"]
                        )

                    except Exception:
                        continue

            retry_count += 1
        else:
            draft["status"] = "NEEDS_REVIEW"
            break

    print("DOC TYPE:", registry_doc["internal_type"])

    print(f"[FINAL] Draft status: {draft['status']} | "
          f"Sections: {len(draft['sections'])}")

    return draft

app.py
import streamlit as st
import requests
from backend.utils.schema_merger import merge_input_groups

API_BASE_URL = "http://127.0.0.1:8000"

st.set_page_config(
    page_title="DocForge Hub",
    layout="wide"
)

st.title("DocForge Hub")
st.caption("Generate, Review, Approve & Publish Documents")

st.divider()

# ---------------- SESSION STATE ----------------

if "selected_draft_id" not in st.session_state:
    st.session_state.selected_draft_id = None

if "last_generated_id" not in st.session_state:
    st.session_state.last_generated_id = None

if "generation_in_progress" not in st.session_state:
    st.session_state.generation_in_progress = False

# ---------------- FIELD RENDERING ----------------

def render_field(label, field, key):
    field_type = field["type"]
    if field_type == "text":
        return st.text_input(label, key=key)
    elif field_type == "textarea":
        return st.text_area(label, key=key, height=100)
    elif field_type == "number":
        return st.number_input(label, key=key)
    elif field_type == "boolean":
        return st.checkbox(label, key=key)
    elif field_type == "date":
        return st.date_input(label, key=key)
    elif field_type == "dropdown":
        options = field.get("options", [])
        return st.selectbox(label, options, key=key)
    elif field_type == "multiselect":
        options = field.get("options", [])
        return st.multiselect(label, options, key=key)
    else:
        return st.text_input(label, key=key)


# ---------------- DYNAMIC FORM ----------------

def render_dynamic_form(base_groups, doc_groups, document_name):
    user_inputs = {}
    validation_errors = []

    st.markdown("## Document Details")
    st.divider()

    if base_groups:
        st.markdown("### 1. General Information")

        for group in base_groups:
            for field in group["fields"]:
                key = field["key"]
                label = field["label"]

                if field.get("required"):
                    label += " *"

                value = render_field(label, field, key)
                user_inputs[key] = value

                if field.get("required") and not value:
                    validation_errors.append(f"{field['label']} is required.")

        st.divider()

    if doc_groups:
        st.markdown(f"### 2. {document_name} Information")

        for group in doc_groups:
            st.markdown(f"**{group['group_name']}**")

            for field in group["fields"]:
                key = field["key"]
                label = field["label"]

                if field.get("required"):
                    label += " *"

                value = render_field(label, field, key)
                user_inputs[key] = value

                if field.get("required") and not value:
                    validation_errors.append(f"{field['label']} is required.")

            st.divider()

    return user_inputs, validation_errors

# SECTION VALIDATION BADGE HELPER

def section_quality_badge(section_validation: dict) -> str:
    if not section_validation:
        return ""
    if section_validation.get("valid"):
        wc = section_validation.get("word_count", 0)
        return f"   {wc} words"
    else:
        issues = section_validation.get("issues", [])
        return f"    {len(issues)} issue(s)"

# ---------------- STEP 1: DOCUMENT SELECTION ----------------

st.subheader("Select Document")

col1, col2, col3 = st.columns(3)

with col1:
    departments = [
    "HR",
    "IT Operations",
    "Legal",
    "Marketing",
    "Finance & Accounting",
    "Engineering",
    "Quality Assurance",
    "Security & Compliance",
    "Customer Success",
    "Product Management"
    ]

    department = st.selectbox("Department", departments)

response = requests.get(
    f"{API_BASE_URL}/documents/list",
    params={"department": department}
)

documents_meta = response.json() if response.status_code == 200 else []

with col2:
    document_types = sorted(set(doc["internal_type"] for doc in documents_meta))
    selected_type = st.selectbox("Document Type", ["ALL"] + document_types)

with col3:
    if selected_type == "ALL":
        filtered_docs = documents_meta
    else:
        filtered_docs = [
            doc for doc in documents_meta
            if doc["internal_type"] == selected_type
        ]

    if filtered_docs:
        document_label = st.selectbox(
            "Document",
            [doc["document_name"] for doc in filtered_docs]
        )
        document_filename = document_label
    else:
        document_filename = None

st.divider()


# ---------------- STEP 2: COMPANY PROFILE ----------------

with st.expander("Company Profile", expanded=True):
    company_name = st.text_input("Company Name")
    industry = st.text_input("Industry")
    employee_count = st.number_input("Employee Count", min_value=1)
    region = st.text_input("Operating Region")
    compliance = st.text_input("Compliance Frameworks")
    jurisdiction = st.text_input("Jurisdiction")

st.divider()


# ---------------- STEP 3: LOAD DOCUMENT CONFIG ----------------

if document_filename:

    response = requests.post(
        f"{API_BASE_URL}/documents/preview",
        json={
            "department": department,
            "document_filename": document_filename
        }
    )

    if response.status_code != 200:
        st.error("Failed to load document.")
        st.stop()

    doc = response.json()
    merged_groups = merge_input_groups(doc)

    base_groups = []
    doc_groups = []

    for group in merged_groups:
        if group.get("source") == "base":
            base_groups.append(group)
        else:
            doc_groups.append(group)

    user_inputs, validation_errors = render_dynamic_form(
        base_groups,
        doc_groups,
        doc["document_name"]
    )

    st.divider()

    # ---------------- GENERATE BUTTON ----------------

    if st.button("Generate Draft", use_container_width=True):

        if not company_name or not industry or not region or not jurisdiction:
            st.error("Please complete company profile before generating.")
            st.stop()

        if validation_errors:
            for err in validation_errors:
                st.error(err)
            st.stop()

        for key, value in user_inputs.items():
            if hasattr(value, "isoformat"):
                user_inputs[key] = value.isoformat()

        response = requests.post(
            f"{API_BASE_URL}/documents/generate",
            json={
                "department": department.lower(),
                "document_filename": document_filename,
                "company_profile": {
                    "company_name": company_name,
                    "industry": industry,
                    "employee_count": employee_count,
                    "regions": [region],
                    "compliance_frameworks": [compliance],
                    "default_jurisdiction": jurisdiction
                },
                "document_inputs": user_inputs
            }
        )

        if response.status_code == 200:
            result = response.json()
            st.success("Draft Generated Successfully")
            st.session_state.selected_draft_id = result["draft_id"]
        else:
            st.error(f"Generation failed: {response.status_code}")
            st.error(response.text)


# ---------------- DRAFT LIBRARY (ALWAYS VISIBLE) ----------------

st.divider()
st.subheader("Draft Library")

response = requests.get(f"{API_BASE_URL}/documents/drafts")

if response.status_code == 200:
    drafts = response.json()

    if drafts:

    # Keep only latest draft per document_name
        unique_docs = {}

        for draft in drafts:
            name = draft["document_name"]

            if name not in unique_docs:
                unique_docs[name] = draft
            else:
                # Keep the latest version
                if draft["version"] > unique_docs[name]["version"]:
                    unique_docs[name] = draft

        filtered_drafts = list(unique_docs.values())

        for draft in filtered_drafts:
            col1, col2, col3, col4 = st.columns([4, 2, 1, 1])

            with col1:
                st.write(f"**{draft['document_name']}** (v{draft['version']})")

            with col2:
                st.write(draft["status"])

            with col3:
                if st.button("View", key=f"view_{draft['id']}"):
                    st.session_state.selected_draft_id = draft["id"]

            with col4:
                if st.button("Delete", key=f"delete_{draft['id']}"):
                    requests.delete(
                        f"{API_BASE_URL}/documents/draft/{draft['id']}"
                    )
                    st.success("Draft deleted")
                    st.experimental_rerun()
    else:
        st.info("No drafts found.")


# ---------------- SHOW FULL DOCUMENT CONTENT ----------------

if st.session_state.selected_draft_id:

    response = requests.get(
        f"{API_BASE_URL}/documents/draft/{st.session_state.selected_draft_id}"
    )

    if response.status_code == 200:
        draft_detail = response.json()

        st.divider()
        st.subheader("Draft Content")

        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button("Download PDF"):
                st.markdown(
                    f'<a href="{API_BASE_URL}/documents/export/{st.session_state.selected_draft_id}/pdf" target="_blank">Click here to download PDF</a>',
                    unsafe_allow_html=True
                )

        with col2:
            if st.button("Download DOCX"):
                st.markdown(
                    f'<a href="{API_BASE_URL}/documents/export/{st.session_state.selected_draft_id}/docx" target="_blank">Click here to download DOCX</a>',
                    unsafe_allow_html=True
                )

        with col3:
            if st.button("Download XLS"):
                st.markdown(
                    f'<a href="{API_BASE_URL}/documents/export/{st.session_state.selected_draft_id}/xls" target="_blank">Click here to download XLS</a>',
                    unsafe_allow_html=True
                )


        st.divider()
        
        st.subheader("Full Document Preview")

        full_document_text = ""

        for section in draft_detail["sections"]:
            st.markdown(f"## {section['section_name']}")

            if section.get("structured_content"):
                table_data = section["structured_content"]["table"]
                headers = table_data["headers"]
                rows = table_data["rows"]

                import pandas as pd
                df = pd.DataFrame(rows, columns=headers)
                st.table(df)

            else:
                st.markdown(section["content"])

            st.divider()

    else:
        st.error("Failed to load draft")

document.py
from fastapi import APIRouter, HTTPException
from backend.api.schemas import DocumentPreviewRequest
from backend.models.company_profile import CompanyProfile
from backend.api.schemas import CompanyProfileCreate
from backend.registry.db_loader import load_document_from_db
from backend.generation.generator import generate_draft
from backend.api.schemas import DocumentGenerateRequest
from sqlalchemy.orm import Session
from fastapi import Depends
from backend.dependencies import get_db
from backend.db_models import Draft, DraftSection
from backend.db_models import Document
from datetime import datetime, timezone
from fastapi.responses import StreamingResponse
from backend.export.exporter import generate_docx, generate_pdf, generate_xls
from backend.export.docx_formatter import build_docx
import io
from sqlalchemy import func


router = APIRouter(prefix="/documents", tags=["Documents"])


@router.post("/preview")
def preview_document(
    payload: DocumentPreviewRequest,
    db: Session = Depends(get_db)
):
    try:
        doc = load_document_from_db(
            db=db,
            department=payload.department,
            document_filename=payload.document_filename
        )
        return doc

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate")
def generate_document(
    payload: DocumentGenerateRequest,
    db: Session = Depends(get_db)
):
    try:
        registry_doc = load_document_from_db(
            db=db,
            department=payload.department,
            document_filename=payload.document_filename
        )

        draft_result = generate_draft(
            registry_doc=registry_doc,
            department=payload.department,
            document_filename=payload.document_filename,
            company_profile=payload.company_profile,
            document_inputs=payload.document_inputs,
            user_notes=payload.user_notes
        )

        draft = Draft(
            document_name=registry_doc["document_name"],
            department=payload.department,
            status=draft_result["status"],  
            version=1,     
            regeneration_count=draft_result["generation_metadata"].get("retry_count", 0),      
        )
        db.add(draft)
        db.commit()
        db.refresh(draft)

        for idx, section in enumerate(draft_result["sections"], start=1):
            db_section = DraftSection(
                draft_id=draft.id,
                section_name=section["name"],
                section_order=idx,
                content=section["content"],
                structured_content=section.get("structured_content")
            )
            db.add(db_section)

        db.commit()

        return {
            "draft_id": draft.id,
            "status": "draft_saved",
            "message": "Draft generated and stored successfully"
        }

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/drafts")
def list_drafts(db: Session = Depends(get_db)):
    drafts = (
        db.query(Draft)
        .order_by(Draft.created_at.desc())
        .all()
    )

    return [
        {
            "id": d.id,
            "document_name": d.document_name,
            "status": d.status,
            "version": d.version
        }
        for d in drafts
    ]

@router.delete("/draft/{draft_id}")
def delete_draft(draft_id: int, db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id).first()

    if not draft:
     raise HTTPException(status_code=404, detail="Draft not found")


    db.delete(draft)
    db.commit()

    return {"message": "Draft deleted successfully"}

@router.get("/list")
def list_documents(department: str, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(func.lower(Document.department) == department.lower()).all()

    return [
        {
            "document_name": d.document_name,
            "internal_type": d.internal_type
        }
        for d in docs
    ]

@router.get("/draft/{draft_id}")
def get_draft_detail(draft_id: int, db: Session = Depends(get_db)):

    draft = db.query(Draft).filter(Draft.id == draft_id).first()

    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    sections = (
        db.query(DraftSection)
        .filter(DraftSection.draft_id == draft_id)
        .order_by(DraftSection.section_order.asc())
        .all()
    )

    return {
        "id": draft.id,
        "document_name": draft.document_name,
        "status": draft.status,
        "version": draft.version,
        "sections": [
            {
                "section_name": s.section_name,
                "content": s.content,
                "structured_content": s.structured_content
            }
            for s in sections
        ]
    }

@router.post("/regenerate-section")
def regenerate_section(
    draft_id: int,
    section_name: str,
    improvement_note: str,
    db: Session = Depends(get_db)
):
    draft = db.query(Draft).filter(Draft.id == draft_id).first()

    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")

    section = db.query(DraftSection).filter(
        DraftSection.draft_id == draft_id,
        DraftSection.section_name == section_name
    ).first()

    if not section:
        raise HTTPException(status_code=404, detail="Section not found")

    try:
        from backend.generation.generator import regenerate_section_llm

        improved_content = regenerate_section_llm(
            draft={
                "source_document": {
                    "internal_type": draft.document_name,
                    "risk_level": "MEDIUM",
                    "department": draft.department
                }
            },
            section={
                "name": section.section_name,
                "content": section.content
            },
            issues=[improvement_note]
        )

        section.content = improved_content
        section.regeneration_count += 1

        draft.status = "NEEDS_REVIEW"

        db.commit()

        return {"message": "Section regenerated successfully"}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/export/{draft_id}/{file_type}")
def export_draft(draft_id: int, file_type: str, db: Session = Depends(get_db)):

    draft_obj = db.query(Draft).filter(Draft.id == draft_id).first()

    if not draft_obj:
        raise HTTPException(status_code=404, detail="Draft not found")

    sections = (
        db.query(DraftSection)
        .filter(DraftSection.draft_id == draft_id)
        .order_by(DraftSection.section_order.asc())
        .all()
    )

    # Fetch original document metadata
    doc_meta = db.query(Document).filter(
        Document.document_name == draft_obj.document_name,
        Document.department == draft_obj.department
    ).first()

    internal_type = doc_meta.internal_type if doc_meta else ""
    risk_level = doc_meta.risk_level if doc_meta else "MEDIUM"

    draft_dict = {
        "source_document": {
            "document_name": draft_obj.document_name,
            "department": draft_obj.department,
            "internal_type": internal_type,
            "risk_level": risk_level
        },
        "version": f"v{draft_obj.version}",
        "status": draft_obj.status,
        "generation_metadata": {
            "generated_at": draft_obj.created_at.isoformat() if draft_obj.created_at else ""
        },
        "sections": [
            {
                "name": s.section_name,
                "content": s.content,
                "mandatory": True
            }
            for s in sections
        ]
    }

    filename = draft_obj.document_name.replace(" ", "_")

    if file_type == "docx":
        docx_bytes = build_docx(draft_dict)

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'}
        )

    elif file_type == "pdf":
        buffer = generate_pdf(draft_obj)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'}
        )

    elif file_type == "xls":
        buffer = generate_xls(draft_obj)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}.xlsx"'}
        )

    else:
        raise HTTPException(status_code=400, detail="Invalid export type")

@router.post("/company-profile")
def create_company_profile(profile: CompanyProfileCreate, db: Session = Depends(get_db)):
    db_profile = CompanyProfile(**profile.model_dump())
    db.add(db_profile)
    db.commit()
    db.refresh(db_profile)
    return db_profile

docs_formatter.py
import io
import re
from datetime import datetime
from docx.shared import Inches, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(draft: dict) -> bytes:
    # from docx import Document
    # from docx.shared import Pt
    # from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title Page ─────────────────────────────
    meta = draft.get("source_document", {})
    doc_name = meta.get("document_name", "Document")
    department = meta.get("department", "")
    version = draft.get("version", "v1.0")

    title = doc.add_heading(doc_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Department: {department}")
    doc.add_paragraph(f"Version: {version}")

    # ── Sections ───────────────────────────────
    for section in draft.get("sections", []):
        section_name = section.get("name", "")
        content = section.get("content", "")
        structured = section.get("structured_content")

        # Skip completely empty section
        if not content.strip() and not structured:
            continue

        # Section Heading
        doc.add_heading(section_name, level=1)

        # ── If structured table exists ─────────────────
        if structured and "table" in structured:

            table_data = structured["table"]
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])

            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))

            # Header row
            for col, header in enumerate(headers):
                table.rows[0].cells[col].text = str(header)

            # Data rows
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.rows[i + 1].cells[j].text = str(cell)

        # ── Otherwise normal text ──────────────────────
        elif content.strip():

            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        doc.add_page_break()

    # ── Save ───────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

db_models.py
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey, Boolean
from datetime import datetime
from sqlalchemy.dialects.postgresql import JSONB
from sqlalchemy.sql import func
from backend.database import Base
from sqlalchemy.orm import relationship

class Draft(Base):
    __tablename__ = "drafts"

    id = Column(Integer, primary_key=True, index=True)
    document_name = Column(String(255))
    department = Column(String(100))
    status = Column(String(50))
    version = Column(Integer, default=1) 
    created_at = Column(DateTime, server_default=func.now())
    regeneration_count = Column(Integer, default=0)

    sections = relationship("DraftSection", back_populates="draft", cascade="all, delete")


class DraftSection(Base):
    __tablename__ = "draft_sections"

    id = Column(Integer, primary_key=True, index=True)
    draft_id = Column(Integer, ForeignKey("drafts.id", ondelete="CASCADE"))
    section_name = Column(String(255))
    section_order = Column(Integer)
    content = Column(Text)
    structured_content = Column(JSONB, nullable=True)
    created_at = Column(DateTime, server_default=func.now())
    regeneration_count = Column(Integer, default=0)

    draft = relationship("Draft", back_populates="sections") 

class Document(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    document_name = Column(String(255))
    department = Column(String(100))
    internal_type = Column(String(100))
    risk_level = Column(String(50))
    approval_required = Column(Boolean)
    versioning_strategy = Column(String(50))
    regeneration_count = Column(Integer, default=0)
    # last_generated_at = Column(DateTime)  
    version = Column(Integer, default=1)
    created_at = Column(DateTime, server_default=func.now())

    sections = Column(JSONB)
    input_groups = Column(JSONB)


class CompanyProfile(Base):
    __tablename__ = "company_profiles"

    id = Column(Integer, primary_key=True, index=True)

    company_name = Column(String(100), nullable=False)
    industry = Column(String(100), nullable=True)
    employee_count = Column(String, nullable=True)
    regions = Column(String, nullable=True)  # store as comma-separated
    compliance_frameworks = Column(String, nullable=True)
    default_jurisdiction = Column(String, nullable=True)

    company_profile = Column(JSONB, nullable=True)












####################################################
CRITICAL OUTPUT RULES:

You must choose ONE of two formats:

FORMAT A — Plain Text (Default for most sections):
- If this section is NOT explicitly listed in the "WHEN TO USE TABLES" list,
you MUST use plain text format.
Never return table format for any other section.
- Write normal paragraph content
- Use markdown: **bold**, - bullets, 1. numbered lists
- Use ### for subheadings if needed
- DO NOT use JSON
- DO NOT mention tables unless explicitly needed
- DO NOT return code blocks
- Just write the actual content

FORMAT B — Table (Only if section explicitly requires tabular data):
- Return ONLY this exact JSON structure:
{
  "table": {
    "headers": ["Column1", "Column2", "Column3"],
    "rows": [
      ["Row1Col1", "Row1Col2", "Row1Col3"],
      ["Row2Col1", "Row2Col2", "Row2Col3"]
    ]
  }
}
- NO text before or after the JSON
- NO ```json code fences
- NO explanation
- Headers must be a list of strings
- Each row must be a list matching header count

WHEN TO USE TABLES:
Use table format for:
- "Roles and Responsibilities"
- "Compliance Matrix"  
- "Approval Matrix"
- "Escalation Path"
- Any section where data is inherently tabular

DO NOT use table format for:
- "Purpose"
- "Scope" 
- "Overview"
- "Background"
- Policy statements
- Narrative content

INVALID OUTPUTS (These cause errors):
Returning just the word "table"
Returning "null" or empty string
Returning JSON without proper structure
Mixing text and JSON in same response
Using markdown tables (| col | col |) in text format





+====================================================================
generator.py
import uuid
from datetime import datetime, timezone
import os
from dotenv import load_dotenv
from backend.models.company_profile import CompanyProfile
from backend.prompts.loader import build_section_prompt, load_prompt
from backend.generation.validator import validate_draft_llm
from backend.prompts.type_behavior import get_type_behavior, should_generate_toc, get_forbidden_phrases
from backend.prompts.risk_behavior import get_risk_behavior
from backend.prompts.section_rules import get_section_rules, get_section_word_limit
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureChatOpenAI
from backend.generation.llm_provider import get_llm
from langchain_core.messages import SystemMessage, HumanMessage

llm = get_llm()

load_dotenv()


def _should_generate_section(doc_type: str, section_name: str) -> bool:
    """
    Returns False for sections that must be skipped for this doc type.
    Currently gates:
      - Table of Contents / Index  →  only for POLICY, SOP, REPORT, HANDBOOK, STRATEGY, PROPOSAL
    All other sections always return True.
    """
    section_lower = section_name.lower()

    toc_keywords = ["table of contents", "index", "contents page"]
    is_toc_section = any(kw in section_lower for kw in toc_keywords)

    if is_toc_section:
        return should_generate_toc(doc_type)

    return True



# SECTION VALIDATOR
# Checks the LLM output for common quality issues.

def _validate_section_output(
    content: str,
    section_name: str,
    doc_type: str
) -> dict:
    """
    Validates a single generated section.

    Returns:
        {
            "valid": bool,
            "issues": list[str],
            "word_count": int,
            "min_words": int,
            "max_words": int
        }
    """
    issues = []
    word_count = len(content.split())
    min_words, max_words = get_section_word_limit(doc_type, section_name)

    repetitive_phrases = [
        "this section constitutes a binding policy requirement",
        "all employees are subject to this policy from their start date",
        "violations of this policy may result in disciplinary action",
        "this policy is reviewed annually"
    ]

    for phrase in repetitive_phrases:
        if phrase in content.lower():
            issues.append(f"Repetitive boilerplate detected: '{phrase}'")
    
    instruction_phrases = [
    "enter the",
    "provide the",
    "complete all fields",
    "insert the",
    "fill in",
    "specify the",
    "record the following"
    ]

    for phrase in instruction_phrases:
        if phrase in content.lower() and doc_type not in ["FORM", "TEMPLATE"]:
            issues.append(f"Instructional language not allowed in {doc_type}: '{phrase}'")


    #  Word count checks 
    if word_count < min_words:
        issues.append(
            f"Too short: {word_count} words (minimum required: {min_words})"
        )

    if word_count > max_words:
        issues.append(
            f"Too long: {word_count} words (maximum allowed: {max_words})"
        )

    #  Section heading repeated inside content body 
    if section_name.lower() in content[:200].lower():
        issues.append(
            "Section heading is repeated inside the content body. "
            "Heading is added externally — remove it from the content."
        )

    #  Placeholder / unfilled text 
    bad_placeholders = [
        "[TO BE FILLED]", "[INSERT HERE]", "[TBD]",
        "TODO:", "[PLACEHOLDER]", "[ADD CONTENT HERE]",
        "[COMPANY NAME]", "[DATE]"
    ]
    for ph in bad_placeholders:
        if ph.upper() in content.upper():
            issues.append(f"Unfilled placeholder found: '{ph}'")

    #  Model preamble leak 
    preamble_phrases = [
        "here is the content",
        "here is the section",
        "below is the content",
        "i'll now generate",
        "the following section",
        "as requested, here"
    ]
    content_start = content[:120].lower()
    for phrase in preamble_phrases:
        if phrase in content_start:
            issues.append(
                f"Model added preamble text: '{phrase}'. "
                "Output must start directly with document content."
            )

    #  Forbidden phrases 
    forbidden = get_forbidden_phrases(doc_type)
    for phrase in forbidden:
        if phrase.lower() in content.lower():
            issues.append(f"Forbidden phrase detected: '{phrase}'")

    #  Empty output 
    if not content.strip():
        issues.append("Generated content is empty.")

    return {
        "valid": len(issues) == 0,
        "issues": issues,
        "word_count": word_count,
        "min_words": min_words,
        "max_words": max_words
    }


# SINGLE SECTION GENERATOR
# Calls AzureOpenAI for one section and validates output.

def _generate_single_section(
    section_name: str,
    mandatory: bool,
    registry_doc: dict,
    company_block: str,
    company_profile: dict,
    inputs_block: str,
    industry_context: str,
    user_notes: str,
    all_sections: list,
    retry: bool = False,
    previous_issues: list = None
) -> dict:
    """
    Generates content for one section.
    Validates the output and returns the full section result.

    Returns:
        {
            "name": str,
            "mandatory": bool,
            "content": str,
            "section_validation": dict   ← NEW: per-section quality result
        }
    """
    doc_type      = registry_doc["internal_type"]
    risk_level    = registry_doc["risk_level"]
    type_behavior_data  = get_type_behavior(doc_type)
    tone = type_behavior_data.get("tone", "professional")
    voice = type_behavior_data.get("voice", "third-person")
    format_style = type_behavior_data.get("format", "")
    rules = type_behavior_data.get("rules", "")
    avg_section_words = type_behavior_data.get("avg_section_words", "")
    risk_behavior  = get_risk_behavior(risk_level)
    section_rules  = get_section_rules(doc_type, section_name)
    company_name = company_profile.get("company_name", "") if company_profile else ""
    industry = company_profile.get("industry", "") if company_profile else ""
    employee_count = company_profile.get("employee_count", "") if company_profile else ""
    region = ", ".join(company_profile.get("regions", [])) if company_profile else ""
    jurisdiction = company_profile.get("default_jurisdiction", "") if company_profile else ""

    # Build TOC section list string for the prompt
    all_sections_str = "\n".join(
        f"{i+1}. {s['name']}"
        for i, s in enumerate(all_sections)
    )
    min_words, max_words = get_section_word_limit(doc_type, section_name)
    if max_words > 300:
        max_words = 300

    forbidden_phrases = get_forbidden_phrases(doc_type)

    context = {
        "document_name":   registry_doc["document_name"],
        "document_type":   doc_type,
        "risk_level":      risk_level,
        "section_name":    section_name,
        "mandatory":       str(mandatory),
        "company_profile": company_profile,
        "document_inputs": inputs_block,
        "industry_context": industry_context,
        "type_behavior":   rules,
        "tone": tone,
        "voice": voice,
        "format_style": format_style,
        "avg_section_words": avg_section_words,
        "risk_behavior":   risk_behavior,
        "section_rules":   section_rules,
        "all_sections":    all_sections_str,
        "toc_required":    str(should_generate_toc(doc_type)).upper(),
        "min_words": min_words,
        "max_words": max_words,
        "company_name": company_name,
        "industry": industry,
        "employee_count": employee_count,
        "region": region,
        "jurisdiction": jurisdiction,
        "forbidden_phrases": "\n".join(forbidden_phrases)
    }

    base_prompt = build_section_prompt(context)

    # ── Add retry context if this is a re-generation ───────
    retry_block = ""
    if retry and previous_issues:
        retry_block = (
            "\n\nPREVIOUS ATTEMPT FAILED VALIDATION — FIX ALL ISSUES BELOW:\n"
            + "\n".join(f"  • {issue}" for issue in previous_issues)
            + "\n\nRe-generate the section addressing every issue listed above.\n"
        )

    full_prompt = f"""
{base_prompt}
{retry_block}
Additional Notes:
{user_notes or "None provided."}
""".strip()

    system_message = f"""
    You are generating the FINAL VERSION of an enterprise {doc_type} document.

    You are NOT:
    - Writing instructions
    - Writing guidance
    - Writing meta commentary
    - Writing a template unless document_type == TEMPLATE
    - Writing placeholders unless document_type == FORM or TEMPLATE

    You ARE:
    - Writing the actual content as it will appear in the published document.

    STRICT LENGTH RULE:
    - Between {min_words} and {max_words} words.
    - Do NOT exceed limit.
    - If exceeded, you FAIL.

    STRICT OUTPUT RULES:
    - Start directly with content.
    - Do NOT repeat section title.
    - Do NOT explain what to do.
    - Do NOT include examples unless explicitly required.
    - Do NOT add filler language.

    SECTION CONTEXT CONTROL:
    - Write content ONLY relevant to the section name.
    - Do NOT introduce topics that belong to other enterprise policies.
    - Do NOT expand scope beyond the purpose of this specific document.

    """

    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=full_prompt)
    ]


    response = llm.invoke(messages)

    try:
        content = response.content.strip()
    except:
        content = str(response).strip()
    
    import json

    structured_content = None

    try:
        parsed = json.loads(content)
        if "table" in parsed:
            structured_content = parsed
            content = ""  # Clear text content if structured table
    except:
        structured_content = None

    max_words_allowed = max_words
    words = content.split()
    if len(words) > max_words_allowed:
        content = " ".join(words[:max_words_allowed])

    # Validate output 
    section_validation = _validate_section_output(
        content=content,
        section_name=section_name,
        doc_type=doc_type
    )
    print("LLM RAW RESULT:", response)
    print("LLM CONTENT:", response.content)

    return {
        "name":               section_name,
        "mandatory":          mandatory,
        "content":            content,
        "structured_content": structured_content,
        "section_validation": section_validation
    }


# SECTION REGENERATION (User-triggered from UI)

def regenerate_section_llm(draft: dict, section: dict, issues: list) -> str:

    template = load_prompt("regeneration_prompt")

    formatted_prompt = template.format(
        document_type=draft["source_document"]["internal_type"],
        risk_level=draft["source_document"]["risk_level"],
        department=draft["source_document"]["department"],
        section_name=section["name"],
        original_content=section["content"],
        issues="\n".join(issues)
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert enterprise document improver."),
        ("human", formatted_prompt)
    ])

    chain = prompt | llm
    result = chain.invoke({})

    return result.content.strip()

def _compress_sections(sections, max_words):
    per_section_budget = max_words // len(sections)

    for s in sections:
        words = s["content"].split()
        if len(words) > per_section_budget:
            s["content"] = " ".join(words[:per_section_budget])

    return sections

# MAIN GENERATE DRAFT

def generate_draft(
    registry_doc: dict,
    department: str,
    document_filename: str,
    company_profile: CompanyProfile = None,
    document_inputs: dict = None,
    user_notes: str = None
) -> dict:
    """
    Generates a full document draft section by section.

    Flow:
      1. Build draft skeleton
      2. Format company profile + user inputs
      3. For each section:
            a. TOC gate  → skip if not needed for this doc type
            b. Generate  → call LLM
            c. Validate  → check word count, placeholders, preamble, etc.
            d. Auto-retry once if validation fails
      4. Run full-draft AI validation with regeneration loop
      5. Return final draft

    Returns: draft dict (same shape as before + section_validation per section)
    """

    #  Step 1: Draft  
    draft = {
        "draft_id": str(uuid.uuid4()),
        "source_document": {
            "department":           department,
            "document_filename":    document_filename,
            "document_name":        registry_doc["document_name"],
            "internal_type":        registry_doc["internal_type"],
            "risk_level":           registry_doc["risk_level"],
            "compliance_alignment": registry_doc.get("compliance_alignment", [])
        },
        "version": "v1.0",
        "status": "DRAFT",
        "generation_metadata": {
            "generated_at":    datetime.now(timezone.utc).isoformat(),
            "generated_by":    "azure_openai",
            "deterministic":   True,
            "prompt_version":  "v2",           # ← bumped to v2 after enhancement
            "toc_generated":   should_generate_toc(registry_doc["internal_type"]),
            "retry_count":     0
        },
        "sections": [],
        "validation": {
            "status": "NOT_RUN",
            "issues": []
        },
        "approval": {
            "required":     registry_doc["approval_required"],
            "approved":     False,
            "approved_by":  None,
            "approved_at":  None
        }
    }

    #  Step 2: Format context blocks 
    company_block = ""
    if company_profile:
        company_block = (
            f"Company Name: {company_profile.get('company_name')}\n"
            f"Industry: {company_profile.get('industry')}\n"
            f"Employee Count: {company_profile.get('employee_count')}\n"
            f"Region: {', '.join(company_profile.get('regions', []))}\n"
            f"Compliance: {', '.join(company_profile.get('compliance_frameworks', []))}\n"
            f"Jurisdiction: {company_profile.get('default_jurisdiction')}\n"
        )

    inputs_block = ""
    if document_inputs:
        for key, value in document_inputs.items():
            inputs_block += f"{key}: {value}\n"


    all_sections     = registry_doc["sections"]

    #  Step 3: Generate each section 
    SECTION_MAX_RETRIES = 1   # One auto-retry per section

    industry_context = load_prompt("industry_context")

    for section in all_sections:
        section_name = section["name"]
        mandatory    = section["mandatory"]

        if section_name.lower() in [
            "security",
            "compliance",
            "data protection",
            "incident response"
        ]: 
            industry_block = industry_context
        else:
            industry_block = ""  


        if not _should_generate_section(
            registry_doc["internal_type"], section_name
        ):
            print(f"[SKIP] '{section_name}' — not required for {registry_doc['internal_type']}")
            continue

        print(f"[GEN]  Generating section: '{section_name}'")

        #  First attempt 
        section_result = _generate_single_section(
            section_name=section_name,
            mandatory=mandatory,
            registry_doc=registry_doc,
            company_profile=company_profile,
            company_block=company_block,
            inputs_block=inputs_block,
            industry_context=industry_block,
            user_notes=user_notes,
            all_sections=all_sections,
            retry  =False
        )

        #  Section-level auto-retry 
        if not section_result["section_validation"]["valid"]:
            issues = section_result["section_validation"]["issues"]
            print(
                f"[WARN] Section '{section_name}' failed validation "
                f"({len(issues)} issue(s)). Retrying..."
            )

            retry_result = _generate_single_section(
                section_name=section_name,
                mandatory=mandatory,
                registry_doc=registry_doc,
                company_block=company_block,
                company_profile=company_profile,    
                inputs_block=inputs_block,
                industry_context=industry_block,
                user_notes=user_notes,
                all_sections=all_sections,
                retry=True,
                previous_issues=issues
            )

            # Use retry result only if it's better (or equal)
            if (
                retry_result["section_validation"]["valid"]
                or len(retry_result["section_validation"]["issues"])
                <= len(issues)
            ):
                section_result = retry_result

        draft["sections"].append(section_result)
        print(
            f"[DONE] '{section_name}' — "
            f"{section_result['section_validation']['word_count']} words | "
            f"valid: {section_result['section_validation']['valid']}"
        )

        #  Global document word cap (max ~2000 words ≈ 4 pages) 

        MAX_TOTAL_WORDS = 1800
        total_words = sum(
            s["section_validation"]["word_count"]
            for s in draft["sections"]
        )

        print(f"[INFO] Total document words before trim: {total_words}")

        if total_words > MAX_TOTAL_WORDS:
            draft["sections"] = _compress_sections(draft["sections"], MAX_TOTAL_WORDS)


    #  Step 4: Full-draft AI validation + regeneration loop 
    MAX_DRAFT_RETRIES = 2
    retry_count       = 0

    validation_result = {"status": "NOT_RUN", "issues": []}

    while retry_count <= MAX_DRAFT_RETRIES:

        try:
            validation_result = validate_draft_llm(draft)
        except Exception as e:
            validation_result = {
                "status": "ERROR",
                "issues": [f"Validation failed: {str(e)}"]
            }

        print(f"[VALIDATE] Status: {validation_result['status']} | Retry: {retry_count}")

        draft["validation"]                            = validation_result
        draft["generation_metadata"]["retry_count"]    = retry_count

        if validation_result["status"] == "PASS":
            draft["status"] = "READY_FOR_APPROVAL"
            break
        else:
            draft["status"] = "NEEDS_REVIEW"

        if retry_count < MAX_DRAFT_RETRIES:
            issues = validation_result.get("issues", [])

            for section in draft["sections"]:
                if section["mandatory"]:
                    try:
                        improved = regenerate_section_llm(
                            draft=draft,
                            section=section,
                            issues=issues
                        )
                        section["content"] = improved

                        # Re-validate the regenerated section
                        section["section_validation"] = _validate_section_output(
                            content=improved,
                            section_name=section["name"],
                            doc_type=registry_doc["internal_type"]
                        )

                    except Exception:
                        continue

            retry_count += 1
        else:
            draft["status"] = "NEEDS_REVIEW"
            break

    print("DOC TYPE:", registry_doc["internal_type"])

    print(f"[FINAL] Draft status: {draft['status']} | "
          f"Sections: {len(draft['sections'])}")

    return draft

17:17 feb 23











########################generator.py
import uuid
from datetime import datetime, timezone
import os
from dotenv import load_dotenv
from backend.models.company_profile import CompanyProfile
from backend.prompts.loader import build_section_prompt, load_prompt
from backend.generation.validator import validate_draft_llm
from backend.prompts.type_behavior import get_type_behavior, should_generate_toc, get_forbidden_phrases
from backend.prompts.risk_behavior import get_risk_behavior
from backend.prompts.section_rules import get_section_rules, get_section_word_limit
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureChatOpenAI
from backend.generation.llm_provider import get_llm
from langchain_core.messages import SystemMessage, HumanMessage
import json
from graphviz import Digraph

llm = get_llm()

load_dotenv()

NEVER_DIAGRAM_SECTIONS = {
        "acknowledgement", "acknowledgement and acceptance",
        "review & revision history", "revision history", "version history",
    }


# Prevent duplicate diagrams inside one draft generation
generated_diagram_signatures = set()


def _should_generate_section(doc_type: str, section_name: str) -> bool:
    """
    Returns False for sections that must be skipped for this doc type.
    Currently gates:
      - Table of Contents / Index  →  only for POLICY, SOP, REPORT, HANDBOOK, STRATEGY, PROPOSAL
    All other sections always return True.
    """
    section_lower = section_name.lower()

    toc_keywords = ["table of contents", "index", "contents page"]
    is_toc_section = any(kw in section_lower for kw in toc_keywords)

    if is_toc_section:
        return should_generate_toc(doc_type)

    return True


# SECTION VALIDATOR
# Checks the LLM output for common quality issues.

def _validate_section_output(
    content: str,
    section_name: str,
    doc_type: str
) -> dict:
    """
    Validates a single generated section.

    Returns:
        {
            "valid": bool,
            "issues": list[str],
            "word_count": int,
            "min_words": int,
            "max_words": int
        }
    """
    issues = []
    if not isinstance(content, str):
        content = str(content)

    word_count = len(content.split())
    min_words, max_words = get_section_word_limit(doc_type, section_name)

    # Structured sections should NOT follow word limits
    if section_name.lower() in [
        "review & revision history",
        "acknowledgement"
    ]:
        min_words = 0
        max_words = 0
    else:
        if max_words > 500:
            max_words = 500

    repetitive_phrases = [
        "this section constitutes a binding policy requirement",
        "all employees are subject to this policy from their start date",
        "violations of this policy may result in disciplinary action",
        "this policy is reviewed annually"
    ]

    for phrase in repetitive_phrases:
        if phrase in content.lower():
            issues.append(f"Repetitive boilerplate detected: '{phrase}'")
    
    instruction_phrases = [
    "enter the",
    "provide the",
    "complete all fields",
    "insert the",
    "fill in",
    "specify the",
    "record the following"
    ]

    for phrase in instruction_phrases:
        if phrase in content.lower() and doc_type not in ["FORM", "TEMPLATE"]:
            issues.append(f"Instructional language not allowed in {doc_type}: '{phrase}'")


    #  Word count checks 
    if word_count < min_words:
        issues.append(
            f"Too short: {word_count} words (minimum required: {min_words})"
        )

    if word_count > max_words:
        issues.append(
            f"Too long: {word_count} words (maximum allowed: {max_words})"
        )

    #  Section heading repeated inside content body 
    if section_name.lower() in content[:200].lower():
        issues.append(
            "Section heading is repeated inside the content body. "
            "Heading is added externally — remove it from the content."
        )

    #  Placeholder / unfilled text 
    bad_placeholders = [
        "[TO BE FILLED]", "[INSERT HERE]", "[TBD]",
        "TODO:", "[PLACEHOLDER]", "[ADD CONTENT HERE]",
        "[COMPANY NAME]", "[DATE]"
    ]
    for ph in bad_placeholders:
        if ph.upper() in content.upper():
            issues.append(f"Unfilled placeholder found: '{ph}'")

    #  Model preamble leak 
    preamble_phrases = [
        "here is the content",
        "here is the section",
        "below is the content",
        "i'll now generate",
        "the following section",
        "as requested, here"
    ]
    content_start = content[:120].lower()
    for phrase in preamble_phrases:
        if phrase in content_start:
            issues.append(
                f"Model added preamble text: '{phrase}'. "
                "Output must start directly with document content."
            )

    #  Forbidden phrases 
    forbidden = get_forbidden_phrases(doc_type)
    for phrase in forbidden:
        if phrase.lower() in content.lower():
            issues.append(f"Forbidden phrase detected: '{phrase}'")

    #  Empty output 
    if not content.strip():
        issues.append("Generated content is empty.")

    return {
        "valid": len(issues) == 0,
        "issues": issues,
        "word_count": word_count,
        "min_words": min_words,
        "max_words": max_words
    }

# def _quick_diagram_trigger(section_name: str, content: str) -> bool:

#     # Only allow diagrams in controlled sections
#     diagram_allowed_sections = [
#         "levels of testing",
#         "process flow",
#         "workflow",
#         "lifecycle",
#         "execution flow",
#         "approval process",
#         "architecture overview"
#     ]
#     name = section_name.lower().strip()

#     return section_name.lower() in diagram_allowed_sections


def _render_flowchart(definition: dict, diagram_type: str = "flowchart") -> str:
    try:
        dot = Digraph(format="png")
        
        # FIXED: higher DPI, tighter size
        dot.attr(dpi="150")      # was 96 — higher = sharper
        dot.attr(margin="0.2")

        nodes = definition.get("nodes", [])
        edges = definition.get("edges", [])

        if diagram_type == "lifecycle":
            dot.attr(rankdir="LR")
            dot.attr(size="6,1.5!")   # wide, short — fits horizontal layout
            dot.attr(nodesep="0.3")
            colors = ["#2E86AB", "#A23B72", "#F18F01", "#C73E1D", "#3B1F2B", "#44BBA4", "#E94F37"]
            for i, node in enumerate(nodes):
                color = colors[i % len(colors)]
                dot.node(str(node), shape="box", style="filled,rounded",
                        fillcolor=color, fontcolor="white",
                        width="1.5", height="0.45", fontsize="9")
            for edge in edges:
                if isinstance(edge, list) and len(edge) == 2:
                    dot.edge(str(edge[0]), str(edge[1]))

        elif diagram_type == "hub":
            dot.attr(size="4,4!")
            center = definition.get("center", nodes[0] if nodes else "Center")
            dot.node(str(center), shape="circle", style="filled",
                    fillcolor="#2C3E50", fontcolor="white",
                    width="1.4", height="1.4", fontsize="10")
            spoke_colors = ["#E74C3C", "#3498DB", "#2ECC71", "#F39C12", "#9B59B6", "#1ABC9C"]
            for i, node in enumerate(nodes):
                color = spoke_colors[i % len(spoke_colors)]
                dot.node(str(node), shape="box", style="filled,rounded",
                        fillcolor=color, fontcolor="white",
                        width="1.2", height="0.4", fontsize="9")
                dot.edge(str(center), str(node), dir="both")

        elif diagram_type == "checklist":
            dot.attr(rankdir="TB", size="3,5!")
            for node in nodes:
                dot.node(str(node), label=f"☐  {node}", shape="box",
                        style="filled", fillcolor="#F8F9FA",
                        fontsize="10", width="2.8", height="0.4")
            for i in range(len(nodes) - 1):
                dot.edge(str(nodes[i]), str(nodes[i+1]), style="invis")

        else:
            # Default flowchart — vertical
            dot.attr(rankdir="TB", size="3,5!")
            dot.attr(nodesep="0.3", ranksep="0.4")
            for node in nodes:
                dot.node(str(node), shape="box", style="filled,rounded",
                        fillcolor="#4A90D9", fontcolor="white",
                        width="2.0", height="0.4", fontsize="9")
            for edge in edges:
                if isinstance(edge, list) and len(edge) == 2:
                    dot.edge(str(edge[0]), str(edge[1]))

        file_id = str(uuid.uuid4())
        folder = os.path.abspath("uploads/diagrams")
        os.makedirs(folder, exist_ok=True)
        path = f"{folder}/{file_id}"
        dot.render(path, cleanup=True)
        return f"{folder}/{file_id}.png"

    except Exception as e:
        print("DIAGRAM RENDER ERROR:", e)
        return None


def _plan_document_diagrams(registry_doc: dict, all_sections: list) -> dict:
    """
    Analyzes the full document globally and decides which sections
    should have diagrams and what type.

    Returns a dict: { "section_name_lowercase": "diagram_type" }
    e.g. {
        "levels of testing": "lifecycle",
        "execution strategy": "flowchart",
        "incident response team": "hub"
    }
    """

    section_list = "\n".join(
        f"{i+1}. {s['name']}"
        for i, s in enumerate(all_sections)
    )

    doc_name = registry_doc["document_name"]
    doc_type = registry_doc["internal_type"]

    prompt = f"""
You are a senior technical writer reviewing an enterprise document structure.

Document Name: {doc_name}
Document Type: {doc_type}

Sections:
{section_list}

YOUR TASK:
Based on the document type and section names, decide which sections 
would genuinely benefit from a visual diagram in a real enterprise document.

Think about what real {doc_type} documents look like — which sections 
typically contain diagrams, flowcharts, lifecycle visuals, or hub diagrams.

Diagram types available:
- "flowchart" → sequential steps/process (e.g. approval flow, deployment steps)
- "lifecycle" → phases/levels in order (e.g. testing pyramid, incident lifecycle, project phases)
- "hub" → central entity with connections (e.g. team structure, system architecture, personas)

Rules:
- Only assign diagrams where they add genuine visual value
- Maximum 3 diagrams per document
- NEVER assign diagrams to: acknowledgement, revision history, 
  summary, appendix, glossary, references, introduction, purpose, 
  scope, objectives, assumptions, dependencies, supporting documents
- Section names that imply a process, hierarchy, team, architecture, 
  timeline, or levels almost always benefit from diagrams

Return STRICT JSON ONLY — a dict mapping lowercase section name to diagram type.
Only include sections that SHOULD have diagrams.

Example:
{{
  "levels of testing": "lifecycle",
  "execution strategy": "flowchart",
  "incident response team": "hub"
}}

If no sections need diagrams:
{{}}
"""

    response = llm.invoke([
        SystemMessage(content="You are an expert technical writer. Return only valid JSON. No markdown."),
        HumanMessage(content=prompt)
    ])

    try:
        raw = response.content.strip()
        if raw.startswith("```"):
            lines = raw.split("\n")
            raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        result = json.loads(raw.strip())
        print(f"[DIAGRAM PLAN] {result}")
        return {k.lower(): v for k, v in result.items()}
    except Exception as e:
        print(f"[DIAGRAM PLAN FAILED] {e}")
        return {}

# SINGLE SECTION GENERATOR
# Calls AzureOpenAI for one section and validates output.

def _generate_single_section(
    section_name: str,
    mandatory: bool,
    registry_doc: dict,
    company_block: str,
    company_profile: dict,
    document_inputs: dict,
    industry_context: str,
    user_notes: str,
    all_sections: list,
    diagram_plan: dict = {},
    retry: bool = False,
    previous_issues: list = None
) -> dict:
    """
    Generates content for one section.
    Validates the output and returns the full section result.

    Returns:
        {
            "name": str,
            "mandatory": bool,
            "content": str,
            "section_validation": dict   ← NEW: per-section quality result
        }
    """

    document_inputs_json = json.dumps(document_inputs or {}, indent=2)

    doc_type      = registry_doc["internal_type"]
    risk_level    = registry_doc["risk_level"]
    type_behavior_data  = get_type_behavior(doc_type)
    tone = type_behavior_data.get("tone", "professional")
    voice = type_behavior_data.get("voice", "third-person")
    format_style = type_behavior_data.get("format", "")
    rules = type_behavior_data.get("rules", "")
    avg_section_words = type_behavior_data.get("avg_section_words", "")
    risk_behavior  = get_risk_behavior(risk_level)
    section_rules  = get_section_rules(doc_type, section_name)
    company_name = company_profile.get("company_name", "") if company_profile else ""
    industry = company_profile.get("industry", "") if company_profile else ""
    employee_count = company_profile.get("employee_count", "") if company_profile else ""
    region = ", ".join(company_profile.get("regions", [])) if company_profile else ""
    jurisdiction = company_profile.get("default_jurisdiction", "") if company_profile else ""
    ceo_name = company_profile.get("ceo_name", "")
    cto_name = company_profile.get("cto_name", "")
    company_background = company_profile.get("company_background", "")

    # Build TOC section list string for the prompt
    all_sections_str = "\n".join(
        f"{i+1}. {s['name']}"
        for i, s in enumerate(all_sections)
    )
    min_words, max_words = get_section_word_limit(doc_type, section_name)
    if max_words > 300:
        max_words = 300

    forbidden_phrases = get_forbidden_phrases(doc_type)

    compliance_frameworks = ", ".join(
        company_profile.get("compliance_frameworks", [])
    ) if company_profile else ""

    context = {
        "document_name":   registry_doc["document_name"],
        "document_type":   doc_type,
        "risk_level":      risk_level,
        "section_name":    section_name,
        "mandatory":       str(mandatory),
        "company_profile": company_profile,
        "document_inputs_json": document_inputs_json,
        "industry_context": industry_context,
        "type_behavior":   rules,
        "tone": tone,
        "voice": voice,
        "format_style": format_style,
        "avg_section_words": avg_section_words,
        "risk_behavior":   risk_behavior,
        "section_rules":   section_rules,
        "all_sections":    all_sections_str,
        "toc_required":    str(should_generate_toc(doc_type)).upper(),
        "min_words": min_words,
        "max_words": max_words,
        "company_name": company_name,
        "industry": industry,
        "employee_count": employee_count,
        "regions": region,
        "default_jurisdiction": jurisdiction,
        "forbidden_phrases": "\n".join(forbidden_phrases),
        "compliance_frameworks": compliance_frameworks,
        "ceo_name": ceo_name,
        "cto_name": cto_name,
        "company_background": company_background
    }

    base_prompt = build_section_prompt(context)

    if section_name.lower() in [
        "acknowledgement",
        "review & revision history"
    ]:
        min_words = 0
        max_words = 200

    if section_name.lower() in [
        "review & revision history",
        "revision history",
        "version history"
    ]:
        base_prompt += """

    OUTPUT STRUCTURE REQUIREMENT:

    Return ONLY ONE JSON table block.

    Do NOT repeat headers.
    Do NOT create multiple table blocks.
    Do NOT duplicate rows.

    Format:

    {
    "type": "table",
    "headers": ["Revision Date", "Version", "Description of Changes", "Approved By"],
    "rows": [
        ["YYYY-MM-DD", "1.0", "Initial creation", "Approver Name"]
    ]
    }
    """

    # STRUCTURE ENFORCEMENT
    if section_name.lower() == "acknowledgement":
        base_prompt += """

    OUTPUT STRUCTURE REQUIREMENT:

    This section MUST be returned as a JSON table block.

    Return ONLY one block in this format:

    {
    "type": "table",
    "headers": ["Field", "Value"],
    "rows": [
        ["Employee Name", ""],
        ["Employee ID", ""],
        ["Designation", ""],
        ["Signature", ""],
        ["Date", ""]
    ]
    }

    Do NOT return paragraph blocks.
    """

    # ── Add retry context if this is a re-generation ───────
    retry_block = ""
    if retry and previous_issues:
        retry_block = (
            "\n\nPREVIOUS ATTEMPT FAILED VALIDATION — FIX ALL ISSUES BELOW:\n"
            + "\n".join(f"  • {issue}" for issue in previous_issues)
            + "\n\nRe-generate the section addressing every issue listed above.\n"
        )

    full_prompt = f"""
{base_prompt}
{retry_block}
Additional Notes:
{user_notes or "None provided."}
""".strip()

    if section_name.lower() in [
        "review & revision history",
        "acknowledgement"
    ]:
        system_message = f"""
    You are generating structured JSON output.

    Return ONLY the JSON table block requested.

    Do NOT:
    - Add narrative text
    - Add explanations
    - Add repeated rows
    - Add multiple table blocks
    - Repeat headers inside rows

    Return exactly ONE clean table.
    """
    else:
        # Look up whether this section should have a diagram
        planned_diagram_type = diagram_plan.get(section_name.lower().strip())

        if planned_diagram_type:
            diagram_instruction = f"""
    DIAGRAM REQUIRED FOR THIS SECTION:
    Include ONE diagram_request block at the end of your response.
    Use diagram_type: "{planned_diagram_type}"
    Extract 3-6 meaningful nodes from your content.
    Nodes must be 2-4 words max.

    Example:
    {{
    "type": "diagram_request",
    "diagram_type": "{planned_diagram_type}",
    "nodes": ["Node 1", "Node 2", "Node 3"],
    "edges": [["Node 1", "Node 2"], ["Node 2", "Node 3"]],
    "center": ""
    }}
    """
        else:
            diagram_instruction = """
    DO NOT include a diagram_request block for this section.
    Return only paragraph blocks.
    """

        system_message = f"""
    You are generating the FINAL VERSION of an enterprise {doc_type} document section.

    STRICT LENGTH RULE:
    - Between {min_words} and {max_words} words.

    OUTPUT FORMAT:
    Return a JSON array of blocks with paragraph content.
    {diagram_instruction}

    STRICT OUTPUT RULES:
    - Start directly with content
    - Do NOT repeat section title
    - Do NOT explain what to do
    - Do NOT add filler language

    Example output WITH diagram:
    [
    {{"type": "paragraph", "content": "Your section content here..."}},
    {{
        "type": "diagram_request",
        "diagram_type": "lifecycle",
        "nodes": ["Stage 1", "Stage 2", "Stage 3"],
        "edges": [["Stage 1", "Stage 2"], ["Stage 2", "Stage 3"]],
        "center": ""
    }}
    ]

    Example output WITHOUT diagram:
    [
    {{"type": "paragraph", "content": "Your section content here..."}}
    ]
    """

    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=full_prompt)
    ]


    response = llm.invoke(messages)

    if section_name.lower() in [
        "review & revision history",
        "revision history",
        "version history"
    ]:
        return {
            "section_name": section_name,
            "mandatory": mandatory,
            "blocks": [
                {
                    "type": "table",
                    "headers": [
                        "Revision Date",
                        "Version",
                        "Description of Changes",
                        "Approved By"
                    ],
                    "rows": [
                        [
                            datetime.now().strftime("%Y-%m-%d"),
                            "1.0",
                            "Initial creation",
                            "Head of Human Resources (HR Director)"
                        ]
                    ]
                }
            ],
            "section_validation": {
                "valid": True,
                "issues": [],
                "word_count": 0,
                "min_words": 0,
                "max_words": 0
            }
        }

    try:
        content = response.content.strip()
    except:
        content = str(response).strip()


    blocks = None
    content_text = ""

    try:
        parsed = json.loads(content)

        # If single object → wrap into list
        if isinstance(parsed, dict):
            parsed = [parsed]

        if isinstance(parsed, list):
            unique_blocks = []
            seen_tables = set()

            for block in parsed:

                if block.get("type") == "table":

                    headers = block.get("headers", [])
                    header_tuple = tuple(h.strip().lower() for h in headers)

                    raw_rows = block.get("rows", [])
                    cleaned_rows = []
                    seen_rows = set()

                    for r in raw_rows:

                        # Normalize row for comparison
                        row_tuple_normalized = tuple(str(cell).strip().lower() for cell in r)

                        # 🚫 Remove header row repeated inside rows
                        if row_tuple_normalized == header_tuple:
                            continue

                        # 🚫 Remove duplicates
                        if row_tuple_normalized in seen_rows:
                            continue

                        seen_rows.add(row_tuple_normalized)
                        cleaned_rows.append(r)

                    block["rows"] = cleaned_rows
                    unique_blocks.append(block)

                else:
                    unique_blocks.append(block)

            blocks = unique_blocks
            print("PARSED BLOCKS COUNT:", len(parsed))
        else:
            raise ValueError("Invalid JSON structure")

    # CHANGE TO — strip markdown fences before parsing:
    except Exception:
        # Try stripping markdown fences first
        cleaned = content.strip()
        if cleaned.startswith("```"):
            lines = cleaned.split("\n")
            cleaned = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        
        try:
            parsed_retry = json.loads(cleaned)
            if isinstance(parsed_retry, dict):
                parsed_retry = [parsed_retry]
            if isinstance(parsed_retry, list):
                blocks = parsed_retry
            else:
                raise ValueError("Still invalid")
        except Exception:
            blocks = [
                {
                    "type": "paragraph",
                    "content": content.strip()
                }
            ]
    combined_text = ""

    for block in blocks:
        if isinstance(block, dict) and block.get("type") == "paragraph":
            combined_text += block.get("content", "") + " "

    content = combined_text.strip()

    # ---------------- DIAGRAM DETECTION ----------------
    # Extract diagram_request block if LLM included one
    final_blocks = []
    diagram_request = None

    for block in blocks:
        if block.get("type") == "diagram_request":
            diagram_request = block
        else:
            final_blocks.append(block)

    blocks = final_blocks

    # Render diagram if requested and valid
    if (
        diagram_request
        and section_name.lower().strip() not in NEVER_DIAGRAM_SECTIONS
        and len(diagram_request.get("nodes", [])) >= 3
    ):
        signature = tuple(diagram_request["nodes"])

        if signature not in generated_diagram_signatures:
            generated_diagram_signatures.add(signature)
            dtype = diagram_request.get("diagram_type", "flowchart")
            diagram_def = {
                "nodes": diagram_request["nodes"],
                "edges": diagram_request.get("edges", []),
                "center": diagram_request.get("center", "")
            }
            image_path = _render_flowchart(diagram_def, dtype)
            file_id = os.path.basename(image_path).replace(".png", "") if image_path else None

            if image_path:
                blocks.append({
                    "type": "diagram",
                    "diagram_type": dtype,
                    "definition": diagram_def,
                    "render_path": image_path,
                    "diagram_url": f"/diagrams/{file_id}.png" if file_id else None,
                    "source": "generated"
                })
                print(f"[DIAGRAM] '{dtype}' generated for '{section_name}'")
    else:
        print(f"[DIAGRAM SKIP] '{section_name}'")

    if max_words > 0:
        max_words_allowed = max_words
        words = content.split()

        if len(words) > max_words_allowed:
            trimmed = " ".join(words[:max_words_allowed])

            # finish sentence if cut mid-way
            if not trimmed.endswith((".", "!", "?")):
                last_period = trimmed.rfind(".")
                if last_period > 50:
                    trimmed = trimmed[: last_period + 1]

            # Only update paragraph blocks, keep diagrams
            new_blocks = []

            for block in blocks:
                if block.get("type") == "paragraph":
                    new_blocks.append({
                        "type": "paragraph",
                        "content": trimmed
                    })
                else:
                    new_blocks.append(block)

            blocks = new_blocks
            content = trimmed

    if section_name.lower() in [
        "review & revision history",
        "acknowledgement"
    ]:
        section_validation = {
            "valid": True,
            "issues": [],
            "word_count": 0,
            "min_words": 0,
            "max_words": 0
        }
    else:
        section_validation = _validate_section_output(
            content=content,
            section_name=section_name,
            doc_type=doc_type
        )
    print("SECTION NAME:", section_name.lower())
    print("LLM RAW RESULT:", response)
    print("LLM CONTENT:", response.content)
    print("FINAL BLOCKS:", blocks)

    if 'parsed' in locals():
        print("PARSED BLOCKS COUNT:", len(parsed))
    else:
        print("PARSED BLOCKS COUNT: JSON parsing failed")

    return {
        "name": section_name,
        "mandatory": mandatory,
        "blocks": blocks,
        "section_validation": section_validation
    }


# SECTION REGENERATION (User-triggered from UI)

def regenerate_section_llm(draft: dict, section: dict, issues: list) -> str:

    template = load_prompt("regenerate_prompt")

    formatted_prompt = template.format(
        document_type=draft["source_document"]["internal_type"],
        risk_level=draft["source_document"]["risk_level"],
        department=draft["source_document"]["department"],
        section_name=section["name"],
        original_content=" ".join(
            block.get("content", "")
            for block in section.get("blocks", [])
            if isinstance(block, dict) and block.get("type") == "paragraph"
        ),
        issues="\n".join(issues)
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert enterprise document improver."),
        ("human", formatted_prompt)
    ])

    chain = prompt | llm
    result = chain.invoke({})

    try:
        parsed = json.loads(result.content)

        if isinstance(parsed, dict):
            parsed = [parsed]

        if isinstance(parsed, list):
            return parsed

    except:
        pass

    # fallback to paragraph
    return [
        {
            "type": "paragraph",
            "content": result.content.strip()
        }
    ]

def _compress_sections(sections, max_words):
    per_section_budget = max_words // len(sections)

    for s in sections:
        combined = " ".join(
            block["content"]
            for block in s["blocks"]
            if block.get("type") == "paragraph"
        )

        words = combined.split()

        if len(words) > per_section_budget:
            trimmed = " ".join(words[:per_section_budget])

            # Keep non-paragraph blocks (diagrams, tables), replace only paragraphs
            non_paragraph_blocks = [
                block for block in s["blocks"]
                if block.get("type") != "paragraph"
            ]

            s["blocks"] = [
                {
                    "type": "paragraph",
                    "content": trimmed
                }
            ] + non_paragraph_blocks  # ← diagrams preserved

            s["section_validation"]["word_count"] = len(trimmed.split())

    return sections

# MAIN GENERATE DRAFT

def generate_draft(
    registry_doc: dict,
    department: str,
    document_filename: str,
    company_profile: CompanyProfile = None,
    document_inputs: dict = None,
    user_notes: str = None
) -> dict:
    """
    Generates a full document draft section by section.

    Flow:
      1. Build draft skeleton
      2. Format company profile + user inputs
      3. For each section:
            a. TOC gate  → skip if not needed for this doc type
            b. Generate  → call LLM
            c. Validate  → check word count, placeholders, preamble, etc.
            d. Auto-retry once if validation fails
      4. Run full-draft AI validation with regeneration loop
      5. Return final draft

    Returns: draft dict (same shape as before + section_validation per section)
    """

    global generated_diagram_signatures
    generated_diagram_signatures = set()

    diagram_plan = _plan_document_diagrams(registry_doc, registry_doc["sections"])

    #  Step 1: Draft  
    draft = {
        "draft_id": str(uuid.uuid4()),
        "source_document": {
            "department":           department,
            "document_filename":    document_filename,
            "document_name":        registry_doc["document_name"],
            "internal_type":        registry_doc["internal_type"],
            "risk_level":           registry_doc["risk_level"],
            "compliance_alignment": registry_doc.get("compliance_alignment", [])
        },
        "version": "v1.0",
        "status": "DRAFT",
        "generation_metadata": {
            "generated_at":    datetime.now(timezone.utc).isoformat(),
            "generated_by":    "azure_openai",
            "deterministic":   True,
            "prompt_version":  "v2",           # ← bumped to v2 after enhancement
            "toc_generated":   should_generate_toc(registry_doc["internal_type"]),
            "retry_count":     0
        },
        "sections": [],
        "validation": {
            "status": "NOT_RUN",
            "issues": []
        },
        "approval": {
            "required":     registry_doc["approval_required"],
            "approved":     False,
            "approved_by":  None,
            "approved_at":  None
        }
    }

    #  Step 2: Format context blocks 
    company_block = ""
    if company_profile:
        company_block = (
            f"Company Name: {company_profile.get('company_name')}\n"
            f"Industry: {company_profile.get('industry')}\n"
            f"Employee Count: {company_profile.get('employee_count')}\n"
            f"Region: {', '.join(company_profile.get('regions', []))}\n"
            f"Compliance: {', '.join(company_profile.get('compliance_frameworks', []))}\n"
            f"Jurisdiction: {company_profile.get('default_jurisdiction')}\n"
        )


    all_sections     = registry_doc["sections"]

    #  Step 3: Generate each section 
    SECTION_MAX_RETRIES = 1   # One auto-retry per section

    industry_context = load_prompt("industry_context")

    for section in all_sections:
        section_name = section["name"]
        mandatory    = section["mandatory"]

        if section_name.lower() in [
            "security",
            "compliance",
            "data protection",
            "incident response"
        ]: 
            industry_block = industry_context
        else:
            industry_block = ""  


        if not _should_generate_section(
            registry_doc["internal_type"], section_name
        ):
            print(f"[SKIP] '{section_name}' — not required for {registry_doc['internal_type']}")
            continue

        print(f"[GEN]  Generating section: '{section_name}'")

        #  First attempt 
        section_result = _generate_single_section(
            section_name=section_name,
            mandatory=mandatory,
            registry_doc=registry_doc,
            company_profile=company_profile,
            company_block=company_block,
            document_inputs=document_inputs, 
            industry_context=industry_block,
            user_notes=user_notes,
            all_sections=all_sections,
            diagram_plan=diagram_plan,
            retry  =False
        )

        #  Section-level auto-retry 
        if not section_result["section_validation"]["valid"]:
            issues = section_result["section_validation"]["issues"]
            print(
                f"[WARN] Section '{section_name}' failed validation "
                f"({len(issues)} issue(s)). Retrying..."
            )

            retry_result = _generate_single_section(
                section_name=section_name,
                mandatory=mandatory,
                registry_doc=registry_doc,
                company_block=company_block,
                company_profile=company_profile,    
                document_inputs=document_inputs, 
                industry_context=industry_block,
                user_notes=user_notes,
                all_sections=all_sections,
                diagram_plan=diagram_plan, 
                retry=True,
                previous_issues=issues
            )

            # Use retry result only if it's better (or equal)
            if (
                retry_result["section_validation"]["valid"]
                or len(retry_result["section_validation"]["issues"])
                <= len(issues)
            ):
                section_result = retry_result

        draft["sections"].append(section_result)
        print(
            f"[DONE] '{section_name}' — "
            f"{section_result['section_validation']['word_count']} words | "
            f"valid: {section_result['section_validation']['valid']}"
        )

        #  Global document word cap (max ~2000 words ≈ 4 pages) 

        MAX_TOTAL_WORDS = 1800
        total_words = sum(
            s["section_validation"]["word_count"]
            for s in draft["sections"]
        )

        print(f"[INFO] Total document words before trim: {total_words}")

        if total_words > MAX_TOTAL_WORDS:
            draft["sections"] = _compress_sections(draft["sections"], MAX_TOTAL_WORDS)


    #  Step 4: Full-draft AI validation + regeneration loop 
    MAX_DRAFT_RETRIES = 2
    retry_count       = 0

    validation_result = {"status": "NOT_RUN", "issues": []}

    while retry_count <= MAX_DRAFT_RETRIES:

        try:
            validation_result = validate_draft_llm(draft)
        except Exception as e:
            validation_result = {
                "status": "ERROR",
                "issues": [f"Validation failed: {str(e)}"]
            }

        print(f"[VALIDATE] Status: {validation_result['status']} | Retry: {retry_count}")

        draft["validation"] = validation_result
        draft["generation_metadata"]["retry_count"] = retry_count

        if validation_result["status"] == "PASS":
            draft["status"] = "READY_FOR_APPROVAL"
            break

        elif validation_result["status"] == "FAIL":
            draft["status"] = "NEEDS_REVIEW"

        else:
            # If ERROR, stop retry loop completely
            draft["status"] = "NEEDS_REVIEW"
            print("Stopping retries due to validation ERROR")
            break

        if retry_count < MAX_DRAFT_RETRIES:
            issues = validation_result.get("issues", [])

            for section in draft["sections"]:
                if section["mandatory"]:
                    improved = regenerate_section_llm(
                        draft=draft,
                        section=section,
                        issues=issues
                    )

                    section["blocks"] = improved

        retry_count += 1

    print("DOC TYPE:", registry_doc["internal_type"])

    print(f"[FINAL] Draft status: {draft['status']} | "
          f"Sections: {len(draft['sections'])}")

    print("TOTAL SECTIONS GENERATED:", len(draft["sections"]))
    return draft
