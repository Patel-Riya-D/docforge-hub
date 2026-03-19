"""
docx_formatter.py

This module is responsible for generating DOCX documents from structured
draft data in the DocForge Hub system.

It converts a JSON-like draft structure into a formatted Microsoft Word document,
including:
- Title page with metadata
- Section headings
- Paragraph content
- Tables
- Diagrams (local or URL-based images)

The output is returned as a byte stream for download via API.
"""
import io
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def build_docx(draft: dict) -> bytes:
    """
    Generate a DOCX document from a structured draft dictionary.

    This function converts a draft object into a formatted Word document.
    It supports multiple content types including paragraphs, tables, and diagrams.

    Workflow:
    1. Create title page with document metadata
    2. Iterate through sections
    3. Render blocks (paragraphs, tables, diagrams)
    4. Export document as byte stream

    Supported Block Types:
        - paragraph: Plain text content
        - table: Structured tabular data
        - diagram: Image (local file or URL)

    Args:
        draft (dict): Structured draft object containing:
            - source_document (dict): Metadata (name, department, company, etc.)
            - version (str): Document version
            - sections (list): List of sections with blocks

    Returns:
        bytes: Binary DOCX file content for download or streaming.

    Raises:
        Exception: Handles internal rendering errors (e.g., image loading issues).

    Example Draft Format:
        {
            "source_document": {
                "document_name": "Security Policy",
                "department": "IT",
                "company_name": "ABC Corp"
            },
            "version": "v1.0",
            "sections": [
                {
                    "name": "Introduction",
                    "blocks": [
                        {"type": "paragraph", "content": "This is intro"},
                        {
                            "type": "table",
                            "headers": ["A", "B"],
                            "rows": [["1", "2"]]
                        }
                    ]
                }
            ]
        }
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title Page ─────────────────────────────
    meta = draft.get("source_document", {})
    doc_name = meta.get("document_name", "Document")
    department = meta.get("department", "")
    version = draft.get("version", "v1.0")

    # --- Title Page ---
    doc.add_paragraph("\n\n\n")

    title = doc.add_heading(doc_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_para.add_run(f"Department: {department}\n")
    meta_para.add_run(f"Version: {version}\n")
    meta_para.add_run(f"Generated On: {datetime.now().strftime('%Y-%m-%d')}\n")

    doc.add_paragraph("\n\n")

    company_name = draft.get("source_document", {}).get("company_name", "")
    if company_name:
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_para.add_run(company_name).bold = True


    doc.add_page_break()

    # ── Sections ───────────────────────────────
    # ── Sections ───────────────────────────────
    for section in draft.get("sections", []):

        section_name = section.get("name", "")
        blocks = section.get("blocks", [])

        heading = doc.add_heading(section_name, level=1)
        heading.runs[0].bold = True

        doc.add_paragraph("")

        # Normalize blocks
        if isinstance(blocks, str):
            blocks = [{"type": "paragraph", "content": blocks}]

        if not isinstance(blocks, list):
            blocks = []

        for block in blocks:

            print("BLOCK:", block)

            # If block is raw string
            if isinstance(block, str):
                doc.add_paragraph(block)
                continue

            if not isinstance(block, dict):
                continue

            block_type = block.get("type")

            # -----------------------------
            # Paragraph
            # -----------------------------
            if block_type == "paragraph":

                text = (
                    block.get("content")
                    or block.get("text")
                    or ""
                )

                if text.strip():
                    doc.add_paragraph(text)

            # -----------------------------
            # Table
            # -----------------------------
            elif block_type == "table":

                headers = block.get("headers", [])
                rows = block.get("rows", [])

                if not headers:
                    continue

                table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = "Table Grid"

                for col_idx, header in enumerate(headers):
                    table.rows[0].cells[col_idx].text = str(header)

                for row_idx, row in enumerate(rows):
                    for col_idx, cell in enumerate(row):
                        table.rows[row_idx + 1].cells[col_idx].text = str(cell)

            # -----------------------------
            # Diagram
            # -----------------------------
            elif block_type == "diagram":

                image_path = block.get("render_path")
                image_url = block.get("diagram_url")

                try:
                    if image_path and os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(5))

                    elif image_url:
                        import requests
                        response = requests.get(image_url)
                        if response.status_code == 200:
                            img = io.BytesIO(response.content)
                            doc.add_picture(img, width=Inches(5))
                        else:
                            doc.add_paragraph("[Diagram not available]")

                    else:
                        doc.add_paragraph("[Diagram missing]")

                except Exception as e:
                    print("IMAGE ERROR:", e)
                    doc.add_paragraph("[Diagram failed to load]")

                doc.add_paragraph("")


    # ── Save ───────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()