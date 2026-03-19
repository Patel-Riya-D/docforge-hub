"""
exporter.py

This module provides utility functions to export generated drafts
into different file formats such as DOCX, PDF, and XLS.

Currently supported:
- DOCX export using python-docx
- (PDF and XLS scaffolding available via reportlab and pandas)

These functions are used by the API layer to generate downloadable
documents from structured draft data.

Note:
This module provides a simpler DOCX generator compared to the
advanced formatter in `docx_formatter.py`.
"""
from docx import Document as DocxDocument
from io import BytesIO
import pandas as pd
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch

def generate_docx(draft):
    """
    Generate a basic DOCX document from draft data.

    This function creates a simple Word document containing:
    - Document title
    - Section headings
    - Paragraph content
    - Tables

    It is a lightweight alternative to the advanced DOCX builder
    (`docx_formatter.py`) and is suitable for quick exports.

    Args:
        draft (dict): Structured draft object containing:
            - source_document (dict): Includes document_name
            - sections (list): List of sections with blocks

    Returns:
        BytesIO: In-memory DOCX file buffer ready for download/streaming.

    Supported Block Types:
        - paragraph: Adds plain text content
        - table: Adds table with headers and rows

    Example:
        draft = {
            "source_document": {"document_name": "Policy"},
            "sections": [
                {
                    "name": "Introduction",
                    "blocks": [
                        {"type": "paragraph", "content": "Intro text"},
                        {
                            "type": "table",
                            "headers": ["A", "B"],
                            "rows": [["1", "2"]]
                        }
                    ]
                }
            ]
        }

    Notes:
        - Does not support advanced styling or diagrams.
        - Uses default Word formatting.
    """
    doc = DocxDocument()

    doc.add_heading(draft["source_document"]["document_name"], level=1)

    for section in draft["sections"]:
        doc.add_heading(section["name"], level=2)

        for block in section["blocks"]:

            if block["type"] == "paragraph":
                doc.add_paragraph(block["content"])

            elif block["type"] == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])

                table = doc.add_table(rows=1, cols=len(headers))

                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = str(header)

                for row in rows:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        row_cells[i].text = str(cell)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer